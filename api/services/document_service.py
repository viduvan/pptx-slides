"""
Document Service — Word and PDF file reading and processing.
Extracted from odin_slides/utils.py for API usage.
"""
import logging
from pathlib import Path
from typing import Callable, Awaitable

from docx import Document

from ..core.config import settings

logger = logging.getLogger("odin_api.document")


def read_docx(file_path: str | Path) -> str:
    """
    Read the full text content of a Word document.

    Args:
        file_path: Path to the .docx file.

    Returns:
        Full text content as a single string.
    """
    try:
        doc = Document(str(file_path))
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        return full_text
    except Exception as e:
        logger.error(f"Error reading Word file: {e}")
        raise


def read_pdf(file_path: str | Path) -> str:
    """
    Read the full text content of a PDF document.

    Args:
        file_path: Path to the .pdf file.

    Returns:
        Full text content as a single string.
    """
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(file_path))
        full_text = ""
        for page in reader.pages:
            text = page.extract_text()
            if text:
                full_text += text + "\n"
        return full_text
    except Exception as e:
        logger.error(f"Error reading PDF file: {e}")
        raise


def read_big_docx(file_path: str | Path, chunk_size: int) -> list[str]:
    """
    Read a large Word document in chunks.

    Args:
        file_path: Path to the .docx file.
        chunk_size: Maximum number of words per chunk.

    Returns:
        List of text chunks.
    """
    try:
        doc = Document(str(file_path))
        chunks = []
        current_text = ""

        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text + "\n"
            if len(current_text.split()) + len(paragraph_text.split()) > chunk_size:
                chunks.append(current_text)
                current_text = paragraph_text
            else:
                current_text += paragraph_text

        if current_text:
            chunks.append(current_text)

        return chunks
    except Exception as e:
        logger.error(f"Error reading big Word file: {e}")
        raise


async def process_document(
    file_path: str | Path,
    summarize_fn: Callable[[str], Awaitable[str]],
) -> tuple[str, bool]:
    """
    Process a Word document: read it, and summarize if it's too large.

    Args:
        file_path: Path to the .docx file.
        summarize_fn: Async function to summarize text chunks.

    Returns:
        Tuple of (processed_content, was_summarized).
    """
    file_ext = Path(file_path).suffix.lower()
    if file_ext == ".pdf":
        word_content = read_pdf(file_path)
    else:
        word_content = read_docx(file_path)
    word_count = len(word_content.split())

    if word_count <= settings.MAX_WORD_COUNT_WITHOUT_SUMMARIZATION:
        return word_content, False

    # Document is large — chunk and summarize
    logger.info(f"Document has {word_count} words, summarizing...")
    chunk_size = word_count // 10
    chunks = read_big_docx(file_path, chunk_size)

    summarized_chunks = []
    for chunk in chunks:
        summary = await summarize_fn(chunk)
        summarized_chunks.append(summary)

    return "\n".join(summarized_chunks), True
